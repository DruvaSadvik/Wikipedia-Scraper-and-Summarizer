import requests
from bs4 import BeautifulSoup
from docx import Document
from vertexai.preview.language_models import TextGenerationModel
import pandas as pd

# Define important keywords related to the topic
important_keywords = [
    "Alexander the Great",
    "Born",
    "Died",
    "Date of Birth",
    "Date of Death",
    "Place of Birth",
    "Place of Death",
    "Parents",
    "Spouse",
    "Children",
    "Early Life",
    "Education",
    "Military Campaigns",
    "Conquests",
    "Battle of Issus",
    "Siege of Tyre",
    "Bucephalus (Horse)",
    "Legacy",
    "Achievements",
    "Religion",
    "Death Cause",
    "Successor",
    "Burial Site",
    "Quotes",
]

# Function to perform the Wikipedia data pipeline
def wikipedia_data_pipeline(topic):
    # Construct the URL for the Wikipedia page
    url = f"https://en.wikipedia.org/wiki/{topic}"

    # Send an HTTP GET request to the URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.text, "html.parser")

    # Create a new Word document
    doc = Document()

    # Finding the main content container of the Wikipedia page
    main_container = soup.find('div', class_='mw-content-container')

    # Create an empty list to store important data
    important_data = []

    # Initialize the Text Generation Model
    text_generation_model = TextGenerationModel.from_pretrained("text-bison@001")

    # Iterate through headings and paragraphs in the Wikipedia content
    for main in main_container.select('h1, h2, h3, h4, h5, h6, p'):
        if main.name in 'h1':
            # If it's an h1 heading, add it as a top-level heading in the document
            doc.add_heading(main.text.strip(), 0)
        elif main.name in 'h2':
            # If it's an h2 heading, add it as a first-level heading in the document
            doc.add_heading(main.text.strip(), level=1)
        elif main.name in ['h3', 'h4', 'h5', 'h6']:
            # If it's an h3, h4, h5, or h6 heading, add it as a lower-level heading in the document
            doc.add_heading(main.text.strip(), level=2)
        else:
            # If it's a paragraph, construct a prompt for text summarization
            prompt = "Summarize this text: " + main.text

            # Generate a summary using the text generation model
            summary = text_generation_model.predict(prompt, max_output_tokens=500).text

            # Add the generated summary as a paragraph in the document
            doc.add_paragraph(summary.strip())

            # Extract important data related to keywords
            for keyword in important_keywords:
                if keyword.lower() in main.text.lower():
                    context = text_generation_model.predict(f'Give a 10-word summary of: {main.text}', max_output_tokens=500).text
                    important_data.append({"Keyword": keyword, "Context": context})

    # Save the generated Word document with the topic as the filename
    doc.save(f'{topic}.docx')

    # Save the extracted important data to an Excel file
    df = pd.DataFrame(important_data)
    df.to_excel(f'{topic}_important_data.xlsx', index=False)

if __name__ == "__main__":
    topic = input("What you want to scrape from Wikipedia? (Try to mention some popular Keyword)")
    wikipedia_data_pipeline(topic)
