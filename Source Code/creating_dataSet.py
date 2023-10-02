import requests
from bs4 import BeautifulSoup
from docx import Document
from vertexai.preview.language_models import TextGenerationModel
import pandas as pd

text_generation_model = TextGenerationModel.from_pretrained("text-bison@001")
doc = Document()
url = "https://en.wikipedia.org/wiki/Alexander_the_Great"
response = requests.get(url)
soup = BeautifulSoup(response.text, "html.parser")
content = {}
main_container = soup.find('div',class_ = 'mw-content-container')

important_keywords = [
    "Alexander the Great",
    "Born",
    "Died",
    "Date of Birth",
    "Date of Death",
    "Place of Birth",
    "Place of Death",
    "Parents",
    "Spouse",
    "Children",
    "Early Life",
    "Education",
    "Military Campaigns",
    "Conquests",
    "Battle of Issus",
    "Siege of Tyre",
    "Bucephalus (Horse)",
    "Legacy",
    "Achievements",
    "Religion",
    "Death Cause",
    "Successor",
    "Burial Site",
    "Quotes",
]

for main in main_container.select('h1,h2,h3,h4,h5,h6,p'): 
    if main.name in 'h1':
        print(main.text)
        doc.add_heading(main.text.strip(),0)
    elif main.name in 'h2':
        print(main.text)
        doc.add_heading(main.text.strip(),level=1)
    elif main.name in ['h3','h4','h5','h6']:
        print(main.text)
        doc.add_heading(main.text.strip(),level=2)
    else:
        prompt = f"Summarize this text, the summary must be coherent and contextually accurate. Provide Best Readable and simple summary {main.text} " 
        summary = text_generation_model.predict(prompt, max_output_tokens=500).text
        doc.add_paragraph(summary.strip())
        important_data = []

   
        for keyword in important_keywords:
            if keyword.lower() in main.text.lower():
                important_data.append({"Keyword": keyword, "Context": text_generation_model.predict(f'give summary in 10 words\n {main.text}', max_output_tokens=500).text})
        df = pd.DataFrame(important_data)
        
        print(summary)
df.to_excel('important_data.xlsx', index=False)
doc.save('Alexander.docx')

